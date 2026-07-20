from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = r"E:\Users\wang\Documents\plan a project\荣耀XH-202631比赛项目实施计划_截至2026-09-01.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "EAF2F8"
PALE = "F4F6F9"
GRAY = "666666"
WHITE = "FFFFFF"
BLACK = "000000"

doc = Document()
sec = doc.sections[0]
sec.page_width = Inches(8.5)
sec.page_height = Inches(11)
sec.top_margin = Inches(0.8)
sec.bottom_margin = Inches(0.75)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.header_distance = Inches(0.35)
sec.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
normal.font.size = Pt(11)
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(8)
normal.paragraph_format.line_spacing = 1.333

for name, size, color, before, after in [
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    st.font.size = Pt(size)
    st.font.bold = True
    st.font.color.rgb = RGBColor.from_string(color)
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)
    st.paragraph_format.keep_with_next = True

for name in ["List Bullet", "List Number"]:
    st = styles[name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
    st.font.size = Pt(11)
    st.paragraph_format.left_indent = Inches(0.375)
    st.paragraph_format.first_line_indent = Inches(-0.194)
    st.paragraph_format.space_after = Pt(4)
    st.paragraph_format.line_spacing = 1.208

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)

def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def set_table_geometry(table, widths_dxa):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for w in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[i]))
            tcW.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_shading(c, PALE)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.1
            if i == 0 and len(str(val)) < 18:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(str(val))
    set_table_geometry(t, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

def add_callout(label, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.08)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_BLUE)
    pPr.append(shd)
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), BLUE)
    pBdr.append(left)
    pPr.append(pBdr)
    r = p.add_run(label + "：")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p.add_run(text)

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.add_run(text)
    return p

def numbered(text):
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)
    return p

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = "PAGE"
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.extend([fldChar1, instrText, fldChar2])

# Running header/footer
hp = sec.header.paragraphs[0]
hp.text = "ResearchSwarm  |  荣耀 XH-202631 比赛项目实施计划"
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hp.runs[0].font.size = Pt(8.5)
hp.runs[0].font.color.rgb = RGBColor.from_string(GRAY)
add_page_number(sec.footer.paragraphs[0])

# Cover: editorial cover pattern
for _ in range(4):
    doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("挑战杯“揭榜挂帅”专项赛 · 项目实施方案")
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor.from_string(BLUE)
p.paragraph_format.space_after = Pt(18)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("ResearchSwarm")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("面向超长程复杂任务的动态异构群体智能引擎")
r.bold = True
r.font.size = Pt(17)
r.font.color.rgb = RGBColor.from_string(BLUE)
p.paragraph_format.space_after = Pt(20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("以 autoresearch-master 为基础的改造与参赛计划")
r.italic = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor.from_string(GRAY)
p.paragraph_format.space_after = Pt(70)

add_table(["赛题编号", "内部冻结日", "官方提交截止"], [["XH-202631", "2026 年 9 月 1 日", "2026 年 9 月 15 日前"]], [2500, 3430, 3430])
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24)
p.add_run("编制日期：2026 年 7 月 17 日").font.color.rgb = RGBColor.from_string(GRAY)
doc.add_page_break()

doc.add_heading("执行摘要", level=1)
doc.add_paragraph("本方案以现有 autoresearch-master 为技术起点，目标不是交付单一的“自动科研”应用，而是将其重构为可复用的动态异构群体智能底座。系统以科研闭环作为主场景，以软件工程闭环作为跨领域验证场景，重点证明超长程记忆保持、动态稀疏拓扑、低熵通信、端边云协同调度以及故障自恢复能力。")
add_callout("关键策略", "9 月 1 日作为功能与材料冻结日，预留 14 天完成长稳测试、视频录制、盖章材料和最终提交。优先保证两个高完成度 Demo 与量化证据链，再投入学习型路由或真实端侧模型切分等冲刺能力。")

doc.add_heading("1. 赛题解读与项目目标", level=1)
doc.add_heading("1.1 赛题硬性要求", level=2)
for x in [
    "在无人工干预下完成多智能体长程协作闭环，并展示中间决策过程和推理轨迹。",
    "通过分布式记忆和上下文压缩/唤醒机制，避免关键事实丢失与全局目标漂移。",
    "使用异构智能体角色和随任务动态变化的稀疏通信拓扑，而非静态全连接。",
    "依据实时性、复杂度和数据敏感等级，在终端、边缘与云端之间自适应选择推理位置。",
    "接受异常、需求变更或节点失效的动态注入，并能自动恢复和完成交付。",
    "至少演示两个高完成度跨领域长程任务，以证明泛化和组合能力。",
]: bullet(x)

doc.add_heading("1.2 总体目标", level=2)
doc.add_paragraph("到 2026 年 9 月 1 日完成 v1.0-rc1：具备可部署、可观测、可恢复、可跨场景验证的群体智能系统；到 9 月 15 日前完成正式材料、演示视频和交付包。建议以约 81/100 为工程目标分，避免在时间有限的情况下平均用力。")

doc.add_heading("1.3 评分反推", level=2)
add_table(
    ["评分项", "分值", "必须提供的证据"],
    [
        ["完整闭环", "15", "端到端轨迹、检查点恢复、失败后自动重规划"],
        ["组织协作", "15", "层级角色、动态组队、每轮稀疏拓扑可视化"],
        ["多任务演示", "10", "科研与软件工程两个跨领域任务"],
        ["应用创新", "25", "效率收益、统一界面、清晰业务价值"],
        ["技术创新", "20", "路由降噪、分层记忆、消融实验和复杂度分析"],
        ["性能效率", "15", "成功率、Token、耗时、兼容性和故障恢复数据"],
    ], [1900, 1000, 6460])

doc.add_heading("2. 产品定位与演示场景", level=1)
doc.add_heading("2.1 项目定位", level=2)
doc.add_paragraph("ResearchSwarm 是一个从高层自然语言意图出发，能够自动组织异构智能体、生成动态协作拓扑、持续维护任务记忆并在异构算力环境中完成闭环交付的自治引擎。autoresearch-master 是第一验证场景和代码基础，但项目架构必须抽象出通用任务状态、Agent 能力注册、消息协议、记忆接口、调度器与可观测系统。")

doc.add_heading("2.2 两个跨领域 Demo", level=2)
add_table(
    ["场景", "端到端链路", "主要评分价值"],
    [
        ["科研闭环（主场景）", "问题定义 → 检索 → 假设 → 实验设计 → 代码/实验 → 结果分析 → 评审 → 报告", "体现超长程、工具调用、记忆保持和复杂推理"],
        ["软件工程（跨域）", "模糊需求 → 规格 → 架构 → 编码 → 测试 → 故障修复 → 文档与交付", "体现跨域泛化、需求变更、节点失效和自治交付"],
    ], [2200, 4400, 2760])

doc.add_heading("3. 总体技术架构", level=1)
doc.add_paragraph("系统由控制平面、异构执行平面、分层记忆平面、端边云资源平面和可观测/评测平面组成。控制平面只保存全局目标、约束、里程碑和验收状态；具体任务按需下沉给不同能力与不同部署位置的执行节点。")
add_table(
    ["平面", "核心组件", "职责"],
    [
        ["控制平面", "Intent Parser、Hierarchical Planner、Dynamic Router、Supervisor、Verifier", "意图解析、DAG 规划、动态组队、质量门禁和局部重规划"],
        ["执行平面", "检索、编码、实验、分析、评审、安全等异构 Agent", "声明能力、成本、延迟、可信度、工具和部署位置并执行子任务"],
        ["记忆平面", "工作记忆、情景记忆、语义记忆、全局黑板", "压缩、唤醒、证据追溯、约束保持和目标漂移检测"],
        ["资源平面", "终端节点、边缘节点、云端节点、Placement Scheduler", "依任务复杂度、时延、隐私、成本和健康状态选址与迁移"],
        ["可观测平面", "Trace、指标、拓扑图、回放、故障注入", "完整展示决策过程并支撑基线、消融、鲁棒性与效率实验"],
    ], [1700, 3500, 4160])

doc.add_heading("4. 端—边—云协同专项设计", level=1)
add_callout("为什么必须突出", "赛题明确要求系统依据子任务实时性和数据敏感等级自动选择推理位置与模型切分，并充分利用云端算力处理高复杂度子任务。该能力不能只停留在部署描述，必须由调度算法、故障迁移和对照实验共同证明。")

doc.add_heading("4.1 三层职责划分", level=2)
add_table(
    ["层级", "适合任务", "模型/组件", "关键约束"],
    [
        ["终端（Device）", "敏感数据预处理、规则检查、轻量意图识别、即时工具执行、离线缓存", "本地小模型、规则引擎、加密缓存、工具代理", "低时延、隐私优先、算力/能耗受限、可能离线"],
        ["边缘（Edge）", "多端请求聚合、检索缓存、上下文压缩、轻量协调、局部记忆和模型网关", "量化模型、向量库、缓存、Edge Supervisor", "区域级算力、网络波动、容量有限、靠近数据源"],
        ["云端（Cloud）", "全局规划、复杂推理、代码生成、大规模检索/实验、跨节点验证", "强模型、全局记忆、Planner、Verifier、实验集群", "能力最强但成本、时延和隐私风险较高"],
    ], [1500, 3500, 2600, 1760])

doc.add_heading("4.2 任务与资源描述", level=2)
doc.add_paragraph("每个子任务携带统一资源描述：复杂度 C、截止时间/实时性 L、数据敏感度 P、预计 Token/算力成本 K、所需能力集合 A、可靠性要求 R。每个节点声明能力、模型质量、当前负载、平均时延、单位成本、隐私等级、健康状态和可用工具。")
doc.add_paragraph("可解释的初版调度评分函数为：")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Score(task, node) = w₁·能力匹配 + w₂·质量收益 + w₃·可靠性 − w₄·时延 − w₅·成本 − w₆·隐私风险 − w₇·迁移开销")
r.bold = True
r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
doc.add_paragraph("调度器先执行硬约束过滤，例如“高敏数据不得离端”“必须在 2 秒内响应”“必须具备代码执行沙箱”，再对候选节点打分。这样既可解释、易调试，也便于形成算法伪代码和复杂度分析。后期如时间允许，再使用 contextual bandit 学习权重。")

doc.add_heading("4.3 调度流程", level=2)
for x in [
    "Planner 生成子任务并标记复杂度、实时性、敏感度、预算、能力和可靠性要求。",
    "Placement Scheduler 读取端、边、云节点的能力、负载、健康状态与历史质量。",
    "先按隐私、截止时间和工具依赖进行硬过滤，再计算候选节点综合分数。",
    "选择主节点和必要的备用节点；只传输最小必要上下文与 artifact 引用。",
    "执行中持续监测延时、Token、错误率和节点健康，触发降级、迁移或重规划。",
    "Verifier 在目标节点或异地节点复核结果，状态写入事件日志并更新调度统计。",
]: numbered(x)

doc.add_heading("4.4 模型选择与切分策略", level=2)
for x in [
    "请求级选择：完整子任务在本地小模型、边缘模型或云端强模型之间切换，是 9 月前最可实现的主方案。",
    "流水线切分：端侧完成脱敏/特征提取，边缘完成检索与压缩，云端只接收最小必要上下文并完成复杂推理。",
    "推测式协同（可选）：端/边小模型生成候选或草稿，云端强模型验证与修正，以降低首 Token 延迟和云端 Token。",
    "神经网络层级切分（Stretch）：真实端侧设备和推理框架具备条件时再验证，不把它设为 MVP 依赖。",
]: bullet(x)

doc.add_heading("4.5 隐私与数据流", level=2)
doc.add_paragraph("端侧先进行数据分类和脱敏，原始高敏数据不离端；边缘只保存区域级缓存、向量索引和带 TTL 的情景记忆；云端接收经过最小化、脱敏和证据引用化的上下文。所有跨层消息带 task_id、敏感等级、允许的 placement、TTL、预算和证据引用，形成可审计的数据流。")

doc.add_heading("4.6 故障迁移与降级路径", level=2)
add_table(
    ["异常", "自动策略", "预期证据"],
    [
        ["云模型超时/限流", "切换备用云模型；或下沉边缘模型完成低风险步骤", "恢复率、额外耗时和质量变化"],
        ["边缘节点失效", "从 checkpoint 在邻近边缘或云端恢复；重新绑定工具", "恢复时间、重复执行次数"],
        ["终端离线", "本地队列和缓存继续可执行步骤；联网后增量同步事件", "离线完成比例、同步一致性"],
        ["网络带宽下降", "压缩上下文、仅传 artifact 引用、降低并发和迁移频率", "传输量、Token、端到端耗时"],
        ["预算骤减", "降低 Top-k、改用小模型、复用缓存、缩减非关键验证", "成本下降及质量损失"],
    ], [1900, 4860, 2600])

doc.add_heading("4.7 无真实端侧设备时的可复现实验方案", level=2)
doc.add_paragraph("比赛前期不应把进度押在真实手机/NPU 适配上。可以在一台或多台计算机上用三个独立服务模拟 Device、Edge、Cloud：Device 使用规则引擎或本地小模型并施加 CPU/内存限制；Edge 使用量化模型、向量库和缓存；Cloud 调用强模型或高算力服务。通过可控网络延时、带宽、节点下线和隐私标签证明调度与迁移机制。材料中明确标注“资源环境模拟”，若后期获得荣耀设备或技术支持，再补真实终端实验。")

doc.add_heading("4.8 端边云实验矩阵与验收指标", level=2)
add_table(
    ["对照方案", "变量", "核心指标"],
    [
        ["全云执行", "所有推理固定在云端", "质量、成本、时延、隐私违规数"],
        ["静态端边云", "按角色固定 placement", "成功率、平均时延、资源利用率"],
        ["自适应调度", "依据任务与节点状态动态选址", "成功率、成本、P95 时延、迁移次数"],
        ["自适应 + 故障注入", "节点失效、超时、带宽下降、预算变化", "恢复率、恢复时间、额外 Token、质量损失"],
    ], [2100, 3500, 3760])
for x in [
    "隐私硬约束违规数 = 0。",
    "相对全云方案，低复杂度任务云端 Token 或调用次数下降至少 25%。",
    "节点故障后自动恢复率不低于 90%，且无需人工修改执行计划。",
    "对实时任务报告 P50/P95 时延，对复杂任务报告质量—成本折中曲线。",
    "UI 能展示每个子任务的 placement、选址理由、迁移轨迹和资源消耗。",
]: bullet(x)

doc.add_heading("5. 动态异构拓扑与低熵通信", level=1)
doc.add_paragraph("Agent 通过 capability registry 声明 capabilities、cost、latency、trust、placement 和 tools。Dynamic Router 根据能力匹配、任务依赖、预期不确定性下降、Token 成本、时延和隐私风险，仅选择 Top-k 接收者形成当前子图。通信拓扑随任务语义、节点状态和预算变化，不预设为静态全连接。")
doc.add_heading("5.1 结构化消息", level=2)
doc.add_paragraph("统一消息至少包含 task_id、sender、receiver_capability、intent、claim、evidence_refs、confidence、requested_action、TTL 和 budget。通过证据引用替代原文重复搬运，并使用 TTL、去重、预算门限和无增益消息过滤抑制噪声级联。")
doc.add_heading("5.2 核心实验", level=2)
doc.add_paragraph("比较单 Agent、静态全连接多 Agent、动态 Top-k 多 Agent三组方案，统一报告成功率、质量分、消息数、总 Token 和耗时。阶段目标是在任务质量不显著下降的情况下，相对静态全连接减少至少 25% 的通信 Token。")

doc.add_heading("6. 分层记忆与超长程连续性", level=1)
add_table(
    ["记忆层", "内容", "生命周期"],
    [
        ["Working Memory", "当前子任务的局部上下文和最近工具结果", "短，完成后压缩"],
        ["Episodic Memory", "事件、决策、错误、恢复、需求变更和结果", "append-only，支持回放"],
        ["Semantic Memory", "事实、约束、摘要、向量索引和产物引用", "跨阶段保留"],
        ["Global Blackboard", "全局目标、里程碑状态、关键决策、验收条件", "全任务强一致"],
    ], [2100, 4700, 2560])
doc.add_paragraph("采用重要度、时效性、依赖中心性和不确定性联合评分进行压缩与唤醒。摘要必须保留来源指针，Verifier 周期性检查关键约束召回和目标漂移。阶段目标为：300 步真实/混合轨迹关键约束召回率不低于 95%，1000 步压力/回放测试无状态损坏。")

doc.add_heading("7. 可恢复长程执行内核", level=1)
for x in [
    "任务持久化为 DAG/状态机；每个节点包含输入、输出、依赖、状态、预算、重试和验收条件。",
    "所有状态变化写入 append-only event log；产物进入 artifact store；按里程碑生成 checkpoint。",
    "执行步骤要求幂等，进程中断后可以从最后一个一致检查点恢复。",
    "Supervisor 处理超时、模型异常、工具异常和需求变更；Verifier 负责结果门禁。",
    "失败时按重试 → 换模型/节点 → 局部重规划 → 全局降级的顺序自动处理。",
]: bullet(x)

doc.add_heading("8. 研发排期与里程碑", level=1)
add_table(
    ["时间", "研发重点", "退出条件"],
    [
        ["7/17-7/20", "代码审计、baseline、两个 Demo 范围冻结", "新环境可复现；指标可采集"],
        ["7/21-7/27", "持久化 DAG、event log、checkpoint、Verifier", "进程中断可恢复；10% 工具失败仍可完成短任务"],
        ["7/28-8/3", "能力注册、动态 Top-k 路由、结构化消息、拓扑 UI", "Token 相对全连接下降 ≥25%，质量不显著下降"],
        ["8/4-8/10", "四层记忆、压缩唤醒、目标漂移检测", "300 步约束召回 ≥95%；千步回放无损坏"],
        ["8/11-8/17", "端边云调度、模型后端、故障注入与迁移", "四类异常自动恢复；至少两类模型后端"],
        ["8/18-8/24", "科研与软件工程 Demo、统一 Web UI、离线回放", "两个 Demo 各连续成功 3 次"],
        ["8/25-8/28", "基线/消融/鲁棒性实验、技术文档和部署包", "所有结论有日志或实验支撑"],
        ["8/29-9/1", "P0/P1 修复、答辩预演、v1.0-rc1 冻结", "空环境部署成功；在线与回放均可演示"],
        ["9/2-9/15", "长稳测试、材料评审、视频、盖章和提交", "9/12 内部截止；最迟 9/14 发送"],
    ], [1650, 4550, 3160])

doc.add_heading("9. MVP、冲刺项与取舍", level=1)
doc.add_heading("9.1 必须完成（MVP）", level=2)
for x in [
    "持久化任务 DAG、事件日志、checkpoint/resume 和幂等执行。",
    "异构角色注册、动态 Top-k 路由、结构化消息和通信指标。",
    "分层记忆、来源追溯、关键约束召回和目标漂移检查。",
    "端边云请求级调度、隐私硬约束、故障迁移和 placement 可视化。",
    "科研与软件工程两个跨域 Demo，在线运行与离线回放双保障。",
    "基线、消融、鲁棒性、Token/时间效率实验及完整技术材料。",
]: bullet(x)
doc.add_heading("9.2 有余力再做（Stretch）", level=2)
for x in [
    "Contextual bandit 或强化学习路由器。",
    "真实荣耀终端/NPU 部署与神经网络层级切分。",
    "基于历史轨迹训练记忆压缩或路由模型。",
    "第三个产业场景或公开高质量长轨迹数据集。",
]: bullet(x)

doc.add_heading("10. 评测体系", level=1)
add_table(
    ["实验", "对照", "指标"],
    [
        ["动态拓扑", "单 Agent / 静态全连接 / 动态 Top-k", "完成率、质量、Token、消息数、耗时"],
        ["分层记忆", "长上下文 / 滚动摘要 / 分层记忆", "约束召回、目标漂移、完成率"],
        ["故障恢复", "无恢复 / 重试 / 重规划+替换", "恢复率、恢复时间、额外成本"],
        ["端边云调度", "全云 / 静态 placement / 自适应", "成本、P95 时延、质量、隐私违规数"],
        ["跨域泛化", "科研 / 软件工程", "端到端完成率、人工干预次数"],
    ], [1900, 3900, 3560])
doc.add_paragraph("统一记录 task_success、quality_score、constraint_recall、tokens_total、messages_total、wall_time、recovery_rate、human_interventions、cost_estimate、placement 和 migration_count。每种关键方案至少运行 5-10 次，报告均值、标准差和失败案例，不只展示最好结果。")

doc.add_heading("11. 风险与应对", level=1)
add_table(
    ["风险", "影响", "应对"],
    [
        ["报名窗口已关闭", "可能失去参赛资格", "立即确认是否已在 6 月 30 日前报名；否则联系赛事方询问补救"],
        ["范围过大", "核心链路无法按时闭环", "MVP 优先；真实端侧切分和学习型路由列为 Stretch"],
        ["原创性不足", "技术创新项失分", "明确上游/自研边界；集中创新路由、记忆、恢复、调度与评测"],
        ["千步表述夸大", "答辩可信度受损", "区分 LLM 步、工具步和回放步，完整披露实验口径"],
        ["现场模型/网络波动", "演示中断", "在线实跑、固定缓存回放、预录视频三层保障"],
        ["缺乏真实端侧设备", "端边云证据不足", "先完成三服务可复现实验，并尽早向比赛专班申请设备/技术支持"],
    ], [2100, 2900, 4360])

doc.add_heading("12. 近期 72 小时行动清单", level=1)
for x in [
    "将 autoresearch-master 放入当前工作区，保留上游地址、许可证和已有修改记录。",
    "提供运行命令、模型/API 依赖、一次成功日志和一次失败日志，完成代码级差距审计。",
    "冻结科研主 Demo 的具体题目和软件工程 Demo 的具体需求及验收条件。",
    "跑通 baseline，保存 Token、耗时、调用链、人工干预次数和最终产物。",
    "确定端、边、云三类模拟节点的模型/工具组合与可用硬件，建立 placement 元数据。",
    "确认报名状态；同步向比赛专班申请端侧设备、模型适配或接口支持。",
]: numbered(x)

doc.add_heading("13. 9 月 1 日交付清单", level=1)
for x in [
    "可部署源码、锁定依赖、配置样例、许可证与第三方清单。",
    "两个跨领域任务的运行脚本、输入、产物、trace 和回放文件。",
    "动态异常、需求变更、节点失效和端边云迁移的注入脚本与结果。",
    "Web 演示界面：DAG、动态拓扑、记忆、placement、迁移、成本和恢复轨迹。",
    "技术文档：架构、路由、记忆、端边云调度、伪代码和复杂度分析。",
    "评测报告：基线、消融、鲁棒性、Token/时间、跨域任务与失败案例。",
    "演示视频、答辩 PPT 初稿、FAQ、申报材料清单和交付包自检脚本。",
]: bullet(x)

add_callout("导师结论", "端边云协同应作为作品的五条主技术线之一，并且要在 UI、故障注入和实验表中持续出现。最稳妥的路线是先完成请求级动态调度和三服务模拟，再把真实终端部署/模型层级切分作为加分项。")

# Prevent rows splitting where possible, set all fonts and metadata.
for table in doc.tables:
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cantSplit = OxmlElement("w:cantSplit")
        trPr.append(cantSplit)
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "等线")
                    run.font.size = Pt(9.5)

doc.core_properties.title = "荣耀 XH-202631 比赛项目实施计划"
doc.core_properties.subject = "ResearchSwarm 动态异构群体智能与端边云协同"
doc.core_properties.author = "参赛团队"
doc.core_properties.keywords = "群体智能, 多智能体, 长程任务, 端边云协同, autoresearch"
doc.save(OUT)
print(OUT)
